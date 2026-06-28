"""
Deposition Splitter Service.

Detects the examination start in legal deposition documents and splits them
into a Cover Section (caption, appearances, stipulations) and an Examination
Section.  Supports PDF, DOCX, and plain-text files.
"""

import os
import re
import uuid
import logging
from dataclasses import dataclass
from typing import Optional

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Result dataclass
# ---------------------------------------------------------------------------

@dataclass
class DepositionSplitResult:
    """Result of a deposition split operation."""
    cover_pdf_path: str          # relative path inside storage/
    examination_pdf_path: str    # relative path inside storage/
    exam_start_page: int         # 1-indexed page where examination begins
    confidence_score: int        # 0–100
    total_pages: int
    cover_page_count: int
    exam_page_count: int


# ---------------------------------------------------------------------------
# Scoring engine — determines whether a page is the examination start
# ---------------------------------------------------------------------------

_EXAM_KEYWORDS = re.compile(
    r"\b(DIRECT\s+EXAMINATION|CROSS\s+EXAMINATION|REDIRECT\s+EXAMINATION"
    r"|EXAMINATION\s+BY|EXAMINATION)\b",
    re.IGNORECASE,
)

_ATTORNEY_PATTERN = re.compile(
    r"\bB\s*Y\s+M[RS]\s*\.|BY\s+M[RS]\.",
    re.IGNORECASE,
)

_Q_PATTERN = re.compile(r"\bQ\.")
_A_PATTERN = re.compile(r"\bA\.")

_SWORN_PATTERN = re.compile(r"having first been duly sworn", re.IGNORECASE)


def score_page(text: str) -> int:
    """Score a page's text for likelihood of being the examination start.

    Rules (cumulative):
        +50  examination keyword present
        +20  attorney pattern (BY MR. / BY MS.)
        +30  Q/A count >= 3
        +15  "having first been duly sworn"
    """
    if not text:
        return 0

    score = 0

    # Rule 1 — Examination keywords
    if _EXAM_KEYWORDS.search(text):
        score += 50

    # Rule 2 — Attorney pattern
    if _ATTORNEY_PATTERN.search(text):
        score += 20

    # Rule 3 — Q/A density
    q_count = len(_Q_PATTERN.findall(text))
    a_count = len(_A_PATTERN.findall(text))
    if q_count >= 3 or a_count >= 3:
        score += 30

    # Rule 4 — Witness sworn
    if _SWORN_PATTERN.search(text):
        score += 15

    return score


# ---------------------------------------------------------------------------
# Text extraction helpers (per-page / per-chunk)
# ---------------------------------------------------------------------------

def _extract_pages_pdf(pdf_path: str) -> list[str]:
    """Return a list of text strings, one per page, from a PDF."""
    doc = fitz.open(pdf_path)
    pages = [doc.load_page(i).get_text("text") for i in range(len(doc))]
    doc.close()
    return pages


def _extract_pages_docx(docx_path: str) -> list[str]:
    """Return a list of text strings grouped into virtual pages from a DOCX.

    DOCX files don't have inherent page boundaries, so we group every 25
    lines (matching typical deposition transcript formatting) into a virtual
    page for scoring purposes.
    """
    from docx import Document

    doc = Document(docx_path)
    all_lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_lines.append(text)

    # Group into virtual pages of ~25 lines
    page_size = 25
    pages: list[str] = []
    for i in range(0, len(all_lines), page_size):
        chunk = all_lines[i : i + page_size]
        pages.append("\n".join(chunk))

    return pages if pages else [""]


def _extract_pages_txt(txt_path: str) -> list[str]:
    """Return a list of text strings grouped into virtual pages from a TXT file.

    Groups every 25 lines into a virtual page.
    """
    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
        all_lines = [line.rstrip() for line in f if line.strip()]

    page_size = 25
    pages: list[str] = []
    for i in range(0, len(all_lines), page_size):
        chunk = all_lines[i : i + page_size]
        pages.append("\n".join(chunk))

    return pages if pages else [""]


def extract_pages(file_path: str) -> list[str]:
    """Extract text grouped by page/virtual-page from any supported file type."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pages_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_pages_docx(file_path)
    elif ext in (".txt", ".text", ".csv", ".json"):
        return _extract_pages_txt(file_path)
    else:
        # Attempt to read as plain text
        try:
            return _extract_pages_txt(file_path)
        except Exception:
            logger.warning(f"Cannot extract pages from unsupported file type: {ext}")
            return []


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------

_CONFIDENCE_THRESHOLD = 70


def detect_examination_start(pages: list[str]) -> tuple[int, int]:
    """Detect the first examination page.

    Args:
        pages: List of page texts (0-indexed internally, returned as 1-indexed).

    Returns:
        (exam_start_page_1indexed, confidence_score)
        If no examination is detected, returns (0, 0).
    """
    for idx, text in enumerate(pages):
        page_score = score_page(text)
        if page_score >= _CONFIDENCE_THRESHOLD:
            logger.info(
                f"Examination detected on page {idx + 1} with confidence {page_score}"
            )
            return idx + 1, page_score

    logger.warning("No examination start detected in document.")
    return 0, 0


# ---------------------------------------------------------------------------
# PDF splitting
# ---------------------------------------------------------------------------

def _split_pdf(
    source_path: str,
    exam_start_page: int,  # 1-indexed
    output_dir: str,
) -> tuple[str, str]:
    """Split a PDF into cover and examination sections.

    Returns:
        (cover_relative_path, exam_relative_path) — relative to storage root.
    """
    doc = fitz.open(source_path)
    total = len(doc)
    split_idx = exam_start_page - 1  # 0-indexed

    uid = uuid.uuid4().hex[:12]

    # Cover section: pages 0 .. split_idx-1
    cover_name = f"{uid}_cover.pdf"
    cover_path = os.path.join(output_dir, cover_name)
    cover_doc = fitz.open()
    if split_idx > 0:
        cover_doc.insert_pdf(doc, from_page=0, to_page=split_idx - 1)
    cover_doc.save(cover_path)
    cover_doc.close()

    # Examination section: pages split_idx .. end
    exam_name = f"{uid}_examination.pdf"
    exam_path = os.path.join(output_dir, exam_name)
    exam_doc = fitz.open()
    exam_doc.insert_pdf(doc, from_page=split_idx, to_page=total - 1)
    exam_doc.save(exam_path)
    exam_doc.close()

    doc.close()

    return (
        f"raw_data/{cover_name}",
        f"raw_data/{exam_name}",
    )


def _split_docx(
    source_path: str,
    exam_start_page: int,  # 1-indexed virtual page
    output_dir: str,
) -> tuple[str, str]:
    """Split a DOCX into cover and examination text files (saved as PDFs via PyMuPDF).

    Since DOCX virtual pages are 25-line groups, we split at the paragraph
    boundary and produce two PDFs.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(source_path)
    all_paragraphs = [p.text for p in doc.paragraphs]

    # Calculate the line index where examination starts
    page_size = 25
    non_empty_lines: list[str] = []
    line_to_para_idx: list[int] = []  # maps non-empty line index → original paragraph index
    for para_idx, text in enumerate(all_paragraphs):
        stripped = text.strip()
        if stripped:
            non_empty_lines.append(stripped)
            line_to_para_idx.append(para_idx)

    split_line = (exam_start_page - 1) * page_size
    if split_line >= len(non_empty_lines):
        split_line = 0

    # Map back to original paragraph index
    if split_line > 0 and split_line < len(line_to_para_idx):
        split_para_idx = line_to_para_idx[split_line]
    else:
        split_para_idx = 0

    uid = uuid.uuid4().hex[:12]

    # Create cover PDF from cover paragraphs
    cover_text = "\n".join(all_paragraphs[:split_para_idx])
    cover_name = f"{uid}_cover.pdf"
    cover_path = os.path.join(output_dir, cover_name)
    _text_to_pdf(cover_text, cover_path)

    # Create examination PDF from remaining paragraphs
    exam_text = "\n".join(all_paragraphs[split_para_idx:])
    exam_name = f"{uid}_examination.pdf"
    exam_path = os.path.join(output_dir, exam_name)
    _text_to_pdf(exam_text, exam_path)

    return (
        f"raw_data/{cover_name}",
        f"raw_data/{exam_name}",
    )


def _split_txt(
    source_path: str,
    exam_start_page: int,  # 1-indexed virtual page
    output_dir: str,
) -> tuple[str, str]:
    """Split a text file into cover and examination sections (saved as PDFs)."""
    with open(source_path, "r", encoding="utf-8", errors="replace") as f:
        all_lines = f.readlines()

    # Non-empty lines for virtual page calculation
    non_empty_indices: list[int] = []
    for i, line in enumerate(all_lines):
        if line.strip():
            non_empty_indices.append(i)

    page_size = 25
    split_non_empty = (exam_start_page - 1) * page_size
    if split_non_empty >= len(non_empty_indices):
        split_non_empty = 0

    # Map to original line index
    if split_non_empty > 0 and split_non_empty < len(non_empty_indices):
        split_line_idx = non_empty_indices[split_non_empty]
    else:
        split_line_idx = 0

    uid = uuid.uuid4().hex[:12]

    cover_text = "".join(all_lines[:split_line_idx])
    cover_name = f"{uid}_cover.pdf"
    cover_path = os.path.join(output_dir, cover_name)
    _text_to_pdf(cover_text, cover_path)

    exam_text = "".join(all_lines[split_line_idx:])
    exam_name = f"{uid}_examination.pdf"
    exam_path = os.path.join(output_dir, exam_name)
    _text_to_pdf(exam_text, exam_path)

    return (
        f"raw_data/{cover_name}",
        f"raw_data/{exam_name}",
    )


def _text_to_pdf(text: str, output_path: str) -> None:
    """Convert plain text to a simple PDF using PyMuPDF."""
    doc = fitz.open()
    width, height = 612, 792  # Letter size
    margin = 72
    font_size = 11
    line_height = font_size * 1.4
    max_y = height - margin

    lines = text.split("\n")

    page = doc.new_page(width=width, height=height)
    y = margin

    for line in lines:
        if y + line_height > max_y:
            page = doc.new_page(width=width, height=height)
            y = margin

        # Wrap long lines
        wrapped = _wrap_text(line, width - 2 * margin, "courier", font_size)
        for wl in wrapped:
            if y + line_height > max_y:
                page = doc.new_page(width=width, height=height)
                y = margin
            page.insert_text(fitz.Point(margin, y), wl, fontsize=font_size, fontname="courier")
            y += line_height

    doc.save(output_path)
    doc.close()


def _wrap_text(text: str, max_width: float, font: str, size: float) -> list[str]:
    """Word-wrap a line of text to fit within max_width."""
    if not text.strip():
        return [""]
    words = text.split()
    wrapped: list[str] = []
    current: list[str] = []
    for word in words:
        test = " ".join(current + [word])
        w = fitz.get_text_length(test, fontname=font, fontsize=size)
        if w <= max_width:
            current.append(word)
        else:
            if current:
                wrapped.append(" ".join(current))
            current = [word]
    if current:
        wrapped.append(" ".join(current))
    return wrapped if wrapped else [""]


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def process_deposition(
    file_path: str,
    storage_path: str,
) -> Optional[DepositionSplitResult]:
    """Detect examination start and split a deposition document.

    Args:
        file_path:    Absolute path to the uploaded document.
        storage_path: Absolute path to the storage root directory.

    Returns:
        DepositionSplitResult if splitting succeeded, None if examination
        could not be detected (confidence too low).
    """
    logger.info(f"Processing deposition: {file_path}")

    # 1. Extract pages
    pages = extract_pages(file_path)
    if not pages:
        logger.warning("No pages extracted from document.")
        return None

    total_pages = len(pages)

    # 2. Detect examination start
    exam_start, confidence = detect_examination_start(pages)
    if exam_start == 0:
        logger.warning(
            f"Examination start not detected (confidence too low). "
            f"Document will not be split."
        )
        return None

    # 3. Split the document
    output_dir = os.path.join(storage_path, "raw_data")
    os.makedirs(output_dir, exist_ok=True)

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        cover_path, exam_path = _split_pdf(file_path, exam_start, output_dir)
    elif ext in (".docx", ".doc"):
        cover_path, exam_path = _split_docx(file_path, exam_start, output_dir)
    else:
        cover_path, exam_path = _split_txt(file_path, exam_start, output_dir)

    cover_page_count = exam_start - 1
    exam_page_count = total_pages - cover_page_count

    result = DepositionSplitResult(
        cover_pdf_path=cover_path,
        examination_pdf_path=exam_path,
        exam_start_page=exam_start,
        confidence_score=confidence,
        total_pages=total_pages,
        cover_page_count=cover_page_count,
        exam_page_count=exam_page_count,
    )

    logger.info(
        f"Deposition split complete: exam starts at page {exam_start}, "
        f"confidence={confidence}, cover={cover_page_count} pages, "
        f"exam={exam_page_count} pages"
    )

    return result


def split_at_page(
    file_path: str,
    exam_start_page: int,  # 1-indexed, user-selected
    storage_path: str,
) -> Optional[DepositionSplitResult]:
    """Split a PDF/DOCX/TXT at a user-specified page number.

    No keyword detection — the user has already identified the page.
    """
    logger.info(f"Manually splitting deposition at page {exam_start_page}: {file_path}")

    pages = extract_pages(file_path)
    if not pages:
        logger.warning("No pages extracted from document.")
        return None

    total_pages = len(pages)

    if exam_start_page < 1 or exam_start_page > total_pages:
        raise ValueError(f"Page {exam_start_page} is out of range (1-{total_pages})")

    output_dir = os.path.join(storage_path, "raw_data")
    os.makedirs(output_dir, exist_ok=True)

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        cover_path, exam_path = _split_pdf(file_path, exam_start_page, output_dir)
    elif ext in (".docx", ".doc"):
        cover_path, exam_path = _split_docx(file_path, exam_start_page, output_dir)
    else:
        cover_path, exam_path = _split_txt(file_path, exam_start_page, output_dir)

    cover_page_count = exam_start_page - 1
    exam_page_count = total_pages - cover_page_count

    result = DepositionSplitResult(
        cover_pdf_path=cover_path,
        examination_pdf_path=exam_path,
        exam_start_page=exam_start_page,
        confidence_score=100,  # user-selected = 100% confidence
        total_pages=total_pages,
        cover_page_count=cover_page_count,
        exam_page_count=exam_page_count,
    )

    logger.info(
        f"Deposition manual split complete: exam starts at page {exam_start_page}, "
        f"cover={cover_page_count} pages, exam={exam_page_count} pages"
    )

    return result

