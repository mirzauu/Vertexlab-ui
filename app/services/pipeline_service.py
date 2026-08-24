"""
Pipeline service: orchestration, status, transcript, and document management.
"""

from uuid import UUID
from datetime import datetime, timezone
import logging
import json

from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

from app.repositories.pipeline_repo import PipelineRepository
from app.repositories.task_repo import TaskRepository
from app.models.pipeline import PipelineRun, PipelineStep, PipelineStatus, StepStatus
from app.models.transcript import Transcript
from app.models.ai_document import AIDocument
from app.models.task import TaskFile, FileType
from app.pipeline.base import PipelineContext
from app.pipeline.orchestrator import PipelineOrchestrator, PIPELINE_STEPS
from app.core.exceptions import NotFoundError, BadRequestError, ConflictError
from app.schemas.pipeline import AIDocumentUpdate, AIDocumentSaveResponse


class PipelineService:
    """Pipeline orchestration and status management."""

    def __init__(self, pipeline_repo: PipelineRepository, task_repo: TaskRepository, db: AsyncSession):
        self.pipeline_repo = pipeline_repo
        self.task_repo = task_repo
        self.db = db

    async def trigger_pipeline(self, task_id: UUID, org_id: UUID) -> PipelineRun:
        """Create a pipeline run and return it. Execution is started via background task."""
        # Verify task exists
        task = await self.task_repo.get_task_in_org(task_id, org_id)
        if not task:
            raise NotFoundError("Task", str(task_id))

        # Check for existing run
        existing = await self.pipeline_repo.get_by_task(task_id)
        if existing and existing.status in (PipelineStatus.QUEUED, PipelineStatus.PROCESSING):
            raise ConflictError("A pipeline is already running for this task")

        # If there was a previous completed/failed run, delete it
        if existing:
            await self.pipeline_repo.delete(existing)

        # Create pipeline run
        pipeline_run = PipelineRun(task_id=task_id)
        pipeline_run = await self.pipeline_repo.create(pipeline_run)

        # Create step records
        for order, step_cls in enumerate(PIPELINE_STEPS, start=1):
            step = step_cls()
            step_record = PipelineStep(
                pipeline_run_id=pipeline_run.id,
                step_name=step.name,
                step_order=order,
            )
            await self.pipeline_repo.create_step(step_record)

        # Reload with steps
        pipeline_run = await self.pipeline_repo.get_by_task(task_id)

        # Commit immediately to ensure background tasks see the new run and avoid race conditions/locks
        await self.db.commit()

        return pipeline_run

    async def execute_pipeline(self, task_id: UUID, org_id: UUID) -> None:
        """
        Execute the pipeline (called from background task).
        This creates a new session context for the background execution.
        """
        from app.db.session import AsyncSessionLocal
        from app.repositories.pipeline_repo import PipelineRepository
        from app.repositories.task_repo import TaskRepository

        async with AsyncSessionLocal() as session:
            # Create fresh repository instances bound to the background session
            pipeline_repo = PipelineRepository(session)
            task_repo = TaskRepository(session)

            pipeline_run = await pipeline_repo.get_by_task(task_id)
            if not pipeline_run:
                return

            # Build context
            task_files = await task_repo.get_files(task_id)
            context = PipelineContext(
                task_id=task_id,
                organization_id=org_id,
            )

            # Populate file paths
            all_raw_paths = []
            examination_paths = []
            for f in task_files:
                if f.file_type == FileType.AUDIO:
                    context.audio_file_path = f.file_path
                elif f.file_type == FileType.RAW_DATA:
                    if f.file_name.endswith("_examination.pdf"):
                        examination_paths.append(f.file_path)
                    else:
                        all_raw_paths.append(f.file_path)
                        
            context.raw_data_file_paths = examination_paths if examination_paths else all_raw_paths

            # Run orchestrator using the background session
            orchestrator = PipelineOrchestrator(session)
            await orchestrator.run(pipeline_run, context)

            # Safety net: persist transcript content if it wasn't already saved
            # by the STT step (it should have been, but this ensures nothing is lost).
            if context.transcript:
                existing_transcript = await pipeline_repo.get_transcript(task_id)
                if not existing_transcript:
                    transcript = Transcript(
                        task_id=task_id,
                        content={"segments": context.transcript.get("segments", [])},
                        language=context.transcript.get("language", "en"),
                        confidence_score=context.transcript.get("confidence"),
                    )
                    await pipeline_repo.save_transcript(transcript)
                elif not existing_transcript.content:
                    # Only update if content is still missing
                    existing_transcript.content = {"segments": context.transcript.get("segments", [])}
                    existing_transcript.language = context.transcript.get("language", "en")
                    existing_transcript.confidence_score = context.transcript.get("confidence")

            # Save document if generated
            if context.generated_document:
                content_data = context.generated_document["content"]
                if context.generated_document.get("corrected_chunks"):
                    content_data = json.dumps(context.generated_document["corrected_chunks"], ensure_ascii=False)

                doc = AIDocument(
                    task_id=task_id,
                    title=context.generated_document["title"],
                    content=content_data,
                    version=context.generated_document.get("version", 1),
                    is_draft=context.generated_document.get("is_draft", True),
                    corrected_chunks=context.generated_document.get("corrected_chunks"),
                )
                await pipeline_repo.save_document(doc)

            await session.commit()

    async def get_status(self, task_id: UUID, org_id: UUID) -> PipelineRun:
        """Get the pipeline run status for a task (eagerly loading task and files)."""
        run = await self.pipeline_repo.get_by_task_in_org(task_id, org_id)
        if not run:
            task = await self.task_repo.get_task_in_org(task_id, org_id)
            if not task:
                raise NotFoundError("Task", str(task_id))
            raise NotFoundError("Pipeline run for this task")

        return run

    async def get_transcript(self, task_id: UUID, org_id: UUID) -> Transcript:
        """Get the transcript for a task."""
        await self.task_repo.get_task_in_org(task_id, org_id)  # Verify access

        transcript = await self.pipeline_repo.get_transcript(task_id)
        if not transcript:
            raise NotFoundError("Transcript")

        return transcript

    async def get_document(self, task_id: UUID, org_id: UUID) -> AIDocument:
        """Get the latest AI document for a task."""
        await self.task_repo.get_task_in_org(task_id, org_id)

        doc = await self.pipeline_repo.get_document(task_id)
        if not doc:
            raise NotFoundError("AI Document")

        return doc

    async def update_document(
        self, task_id: UUID, org_id: UUID, data: AIDocumentUpdate
    ) -> AIDocumentSaveResponse:
        """
        Update the existing draft in-place without version row explosion,
        reducing save network round-trips from 5 down to 1-2.
        """
        from sqlalchemy import select
        from sqlalchemy.orm import undefer
        from app.models.task import Task

        # 1. Fetch current latest document and verify task belongs to org in a single query
        result = await self.db.execute(
            select(AIDocument)
            .join(AIDocument.task)
            .where(AIDocument.task_id == task_id, Task.organization_id == org_id)
            .order_by(AIDocument.version.desc())
            .options(undefer(AIDocument.content), undefer(AIDocument.corrected_chunks))
        )
        current_doc = result.scalars().first()

        content_data = data.content or ""
        if data.corrected_chunks:
            content_data = json.dumps(data.corrected_chunks, ensure_ascii=False)

        if current_doc:
            if not current_doc.is_draft:
                raise BadRequestError("Cannot edit a finalized document")

            # In-place update: prevents DB row bloat and version explosion
            if data.title:
                current_doc.title = data.title
            if data.corrected_chunks is not None:
                current_doc.corrected_chunks = data.corrected_chunks
            current_doc.content = content_data
            current_doc.updated_at = datetime.now(timezone.utc)
            doc_to_return = current_doc
        else:
            # First draft V1 creation — verify task exists in org
            await self.task_repo.get_task_in_org(task_id, org_id)
            new_doc = AIDocument(
                task_id=task_id,
                title=data.title or "AI-Corrected Proof Document",
                content=content_data,
                version=1,
                is_draft=True,
                parent_id=None,
                corrected_chunks=data.corrected_chunks or [],
            )
            self.db.add(new_doc)
            doc_to_return = new_doc

        await self.db.flush()

        # Return lean save response without re-fetching or echoing full chunks
        return AIDocumentSaveResponse(
            status="saved",
            id=doc_to_return.id,
            task_id=task_id,
            version=doc_to_return.version,
            updated_at=doc_to_return.updated_at or datetime.now(timezone.utc),
        )

    async def finalize_document(self, task_id: UUID, org_id: UUID) -> AIDocument:
        """Mark an AI document as final (no longer a draft)."""
        doc = await self.get_document(task_id, org_id)

        if not doc.is_draft:
            raise BadRequestError("Document is already finalized")

        doc.is_draft = False
        await self.db.flush()

        saved_doc = await self.pipeline_repo.get_document(task_id)
        return saved_doc or doc

    async def get_document_versions(
        self, task_id: UUID, org_id: UUID
    ) -> list[AIDocument]:
        """List all versions of the AI document for a task."""
        await self.task_repo.get_task_in_org(task_id, org_id)
        return await self.pipeline_repo.get_all_document_versions(task_id)

    async def get_detailed_results(self, task_id: UUID, org_id: UUID) -> dict:
        """
        Aggregates all pipeline output in a single optimized DB query with eager joins.
        Reduces DB network round-trips from 5 down to 1.
        """
        from datetime import datetime
        from sqlalchemy import select
        from sqlalchemy.orm import selectinload, undefer
        from app.models.pipeline import PipelineRun
        from app.models.transcript import Transcript
        from app.models.task import Task, TaskFile, FileType

        # ── 1. Single query: load Task + files + pipeline_run + steps + transcript ─────
        task_result = await self.db.execute(
            select(Task)
            .where(Task.id == task_id, Task.organization_id == org_id)
            .options(
                selectinload(Task.files),
                selectinload(Task.pipeline_run).selectinload(PipelineRun.steps),
                selectinload(Task.transcript),
            )
        )
        task = task_result.scalar_one_or_none()
        if not task:
            raise NotFoundError("Task", str(task_id))

        pipeline_run = task.pipeline_run
        if not pipeline_run:
            raise NotFoundError("Pipeline run for this task")

        transcript = task.transcript

        # Load ONLY the single latest document version (prevents fetching 100MB+ of old version rows)
        doc_result = await self.db.execute(
            select(AIDocument)
            .where(AIDocument.task_id == task_id)
            .order_by(AIDocument.version.desc())
            .options(undefer(AIDocument.content), undefer(AIDocument.corrected_chunks))
            .limit(1)
        )
        doc = doc_result.scalars().first()

        # ── Build response from already-loaded data (no extra queries) ─────────
        audio_file_path = None
        pdf_raw_data_path = None
        if task.files:
            for f in task.files:
                if f.file_type == FileType.AUDIO:
                    audio_file_path = f.file_path
                elif f.file_type == FileType.RAW_DATA:
                    pdf_raw_data_path = f.file_path

        # Find matching step's metadata_json directly in memory (avoids Query 5)
        summary = {}
        if pipeline_run and pipeline_run.steps:
            for step in pipeline_run.steps:
                if step.step_name == "matching":
                    summary = step.metadata_json or {}
                    break

        transcribed_data = []
        if transcript and transcript.content:
            if isinstance(transcript.content, dict):
                transcribed_data = transcript.content.get("segments", [])
            elif isinstance(transcript.content, list):
                transcribed_data = transcript.content

        return {
            "audio_file_path": audio_file_path,
            "transcribed_data": transcribed_data,
            "pdf_raw_data": [],  # Exclude unused raw steno chunks payload to minimize network size
            "metadata": {
                "audio_source": audio_file_path,
                "raw_data_source": pdf_raw_data_path,
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "version": "8.0"
            },
            "summary": summary,
            "matches": transcript.matches if transcript else [],
            "document": doc
        }

    async def stream_status(self, task_id: UUID, org_id: UUID):
        """Async generator for Server-Sent Events (SSE) realtime status updates."""
        import asyncio
        import json
        from app.schemas.pipeline import PipelineRunRead

        try:
            while True:
                # Rollback any stale transaction to read fresh data from the background worker
                try:
                    await self.db.rollback()
                except Exception:
                    pass

                try:
                    run = await self.get_status(task_id, org_id)
                except NotFoundError:
                    yield {"event": "error", "data": json.dumps({"detail": "Not found"})}
                    break

                run_schema = PipelineRunRead.model_validate(run)
                yield {"event": "status", "data": run_schema.model_dump_json()}

                if run.status in (PipelineStatus.COMPLETED, PipelineStatus.FAILED):
                    break

                await asyncio.sleep(1.5)

        except asyncio.CancelledError:
            # Client disconnected — this is expected behaviour, not an error.
            # Silently exit the generator so SQLAlchemy can close cleanly.
            logger.debug("SSE client disconnected for task %s", task_id)

    async def get_workstation_data(self, task_id: UUID, org_id: UUID) -> dict:
        """
        Combined endpoint payload for the workstation (Review/Edit) page.

        Returns pipeline results AND the resolved audio file metadata in a
        single DB round-trip, replacing the previous two-request pattern
        (pipeline/results + tasks/{id}/files) on the frontend.
        """
        from app.models.task import TaskFile, FileType
        from sqlalchemy import select, or_

        # Re-use the existing detailed results aggregator (already optimised)
        results = await self.get_detailed_results(task_id, org_id)

        # Fetch audio file row: check FileType.AUDIO or audio mime_type/extension fallback
        audio_result = await self.db.execute(
            select(TaskFile)
            .where(
                TaskFile.task_id == task_id,
                or_(
                    TaskFile.file_type == FileType.AUDIO,
                    TaskFile.mime_type.ilike("audio/%"),
                    TaskFile.file_name.ilike("%.mp3"),
                    TaskFile.file_name.ilike("%.wav"),
                    TaskFile.file_name.ilike("%.m4a"),
                    TaskFile.file_name.ilike("%.aac"),
                    TaskFile.file_name.ilike("%.ogg"),
                    TaskFile.file_name.ilike("%.webm"),
                    TaskFile.file_path.ilike("%.mp3"),
                    TaskFile.file_path.ilike("%.wav"),
                    TaskFile.file_path.ilike("%.m4a"),
                )
            )
            .order_by(TaskFile.uploaded_at.desc())
        )
        audio_file = audio_result.scalars().first()

        audio_file_info = None
        if audio_file:
            stream_url = audio_file.cloudinary_url
            if stream_url and "cloudinary.com" in stream_url and stream_url.endswith(".wav"):
                stream_url = stream_url[:-4] + ".mp3"

            audio_file_info = {
                "id": str(audio_file.id),
                "file_path": audio_file.file_path,
                "file_type": audio_file.file_type.value if hasattr(audio_file.file_type, 'value') else str(audio_file.file_type),
                "cloudinary_url": stream_url,
            }

        return {
            "results": results,
            "audio_file": audio_file_info,
        }

    def _split_qa_lines(self, text: str) -> list[str]:
        import re
        if not text:
            return [""]
        # Split on whitespace followed by Q/A indicators (e.g. Q:, A:, Q., A. case-insensitive)
        sub_lines = re.split(r'\s+(?=[qQaA][:\.](?:\s|$))', text)
        return [s.strip() for s in sub_lines if s.strip()]

    async def get_document_pdf(self, task_id: UUID, org_id: UUID) -> bytes:

        """Fetch the latest AI document and generate a beautifully formatted PDF."""
        from sqlalchemy import select
        from sqlalchemy.orm import selectinload
        from app.models.pipeline import PipelineRun

        # 1. Fetch task and check access
        task = await self.task_repo.get_task_in_org(task_id, org_id)
        if not task:
            raise NotFoundError("Task", str(task_id))

        # 2. Fetch the latest AI Document
        doc = await self.pipeline_repo.get_document(task_id)
        if not doc:
            raise NotFoundError("AI Document")

        # 3. Fetch summary stats and layout profile
        run_result = await self.db.execute(
            select(PipelineRun)
            .where(PipelineRun.task_id == task_id)
            .options(selectinload(PipelineRun.steps))
        )
        pipeline_run = run_result.scalar_one_or_none()
        summary = {}
        layout_profile = None
        if pipeline_run:
            for step in pipeline_run.steps:
                if step.step_name == "matching":
                    summary = step.metadata_json or {}
                elif step.step_name == "data_processing":
                    dp_meta = step.metadata_json or {}
                    layout_profile = dp_meta.get("layout_profile") or dp_meta.get("structure_rules", {}).get("layout_profile")

        # 4. Format lines for PDF
        lines = []

        # Add chunks
        chunks = doc.corrected_chunks
        if not chunks:
            # Fallback to transcript matches
            from app.models.transcript import Transcript
            tr_result = await self.db.execute(
                select(Transcript).where(Transcript.task_id == task_id)
            )
            transcript = tr_result.scalar_one_or_none()
            if transcript and transcript.matches:
                chunks = []
                for m in transcript.matches:
                    chunks.append({
                        "raw_chunk_id": m.get("raw_chunk_id"),
                        "original_raw_text": m.get("raw_chunk_text"),
                        "corrected_text": m.get("raw_chunk_text"),
                        "match_status": m.get("match_status", "unknown"),
                        "confidence_score": m.get("confidence_score", 0),
                        "audio_start_time_sec": m.get("audio_start_time_sec"),
                        "audio_end_time_sec": m.get("audio_end_time_sec"),
                        "speakers": m.get("speakers", [])
                    })

        if chunks:
            for chunk in chunks:
                chunk_text = chunk.get("corrected_text") or chunk.get("original_raw_text") or ""
                lines.extend(self._split_qa_lines(chunk_text))
        else:
            # Fallback to plain content
            for line in doc.content.split("\n"):
                lines.extend(self._split_qa_lines(line))

        # 5. Generate AI PDF bytes using dynamic layout profile
        generated_pdf_bytes = self._generate_pdf_from_lines(doc.title, lines, layout_profile=layout_profile)
        
        # 6. Check if there is a cover file and merge it
        import os
        from app.config import settings
        import fitz
        task_files = await self.task_repo.get_files(task_id)
        cover_files = [f for f in task_files if f.file_name.endswith("_cover.pdf")]
        if cover_files:
            cover_path = os.path.join(settings.STORAGE_PATH, cover_files[0].file_path)
            if os.path.exists(cover_path):
                try:
                    cover_doc = fitz.open(cover_path)
                    ai_doc = fitz.open("pdf", generated_pdf_bytes)
                    cover_doc.insert_pdf(ai_doc)
                    combined_bytes = cover_doc.write()
                    cover_doc.close()
                    ai_doc.close()
                    return combined_bytes
                except Exception as e:
                    logger.warning(f"Failed to merge cover PDF: {e}")
                    
        return generated_pdf_bytes

    def _generate_pdf_from_lines(self, title: str, lines: list[str], layout_profile: dict | None = None) -> bytes:
        import fitz
        
        lp = layout_profile or {}
        width = float(lp.get("page_width", 612.0))
        height = float(lp.get("page_height", 792.0))
        
        font_name = str(lp.get("font_name", "courier"))
        font_size = float(lp.get("font_size", 10.0))
        line_number_font_size = float(lp.get("line_number_font_size", font_size))
        
        top_margin = float(lp.get("top_margin", 95.0))
        bottom_margin = float(lp.get("bottom_margin", 72.0))
        left_margin = float(lp.get("left_margin", 79.0))
        right_margin = float(lp.get("right_margin", 79.0))
        
        max_lines_per_page = int(lp.get("max_lines_per_page", 25))
        
        # Line height from profile or dynamic calculation
        if "line_height" in lp and float(lp["line_height"]) > 0:
            line_height = float(lp["line_height"])
        else:
            available_height = height - top_margin - bottom_margin
            line_height = available_height / max_lines_per_page
            
        line_number_x = float(lp.get("line_number_x", left_margin))
        line_number_width = float(lp.get("line_number_width", 22.0))
        text_start_x = float(lp.get("text_start_x", line_number_x + line_number_width + 12.0))
        
        q_start_x = float(lp.get("q_start_x", text_start_x))
        a_start_x = float(lp.get("a_start_x", text_start_x))
        speaker_start_x = float(lp.get("speaker_start_x", text_start_x))
        
        max_width = width - text_start_x - right_margin
        
        page_num_cfg = lp.get("page_number", {})
        page_num_x = float(page_num_cfg.get("x", width - right_margin - 10.0))
        page_num_y = float(page_num_cfg.get("y", max(25.0, top_margin - 20.0)))
        
        vertical_lines = lp.get("vertical_lines", [])
        
        doc = fitz.open()
        page_num = 1
        
        def start_new_page():
            nonlocal page_num
            p = doc.new_page(width=width, height=height)
            
            # Page Number
            p_num_str = str(page_num)
            p.insert_text(fitz.Point(page_num_x, page_num_y), p_num_str, fontsize=font_size, fontname=font_name)
            
            # Draw line numbers down the margin
            y_line = top_margin
            for i in range(1, max_lines_per_page + 1):
                num_str = str(i)
                nw = fitz.get_text_length(num_str, fontname=font_name, fontsize=line_number_font_size)
                x_offset = line_number_x + line_number_width - nw
                p.insert_text(fitz.Point(x_offset, y_line), num_str, fontsize=line_number_font_size, fontname=font_name)
                y_line += line_height
            
            # Draw vertical lines
            if vertical_lines:
                shape = p.new_shape()
                for vl in vertical_lines:
                    vx = float(vl.get("x", 0))
                    vtop = float(vl.get("top", top_margin - 10.0))
                    vbot = float(vl.get("bottom", height - bottom_margin))
                    shape.draw_line(fitz.Point(vx, vtop), fitz.Point(vx, vbot))
                shape.finish(color=(0, 0, 0), width=0.5)
                shape.commit()
                
            page_num += 1
            return p

        page = start_new_page()
        current_line_on_page = 1
        y = top_margin

        def get_wrapped_lines(text: str, max_w: float) -> list[str]:
            words = text.split()
            if not words:
                return ['']
            wrapped = []
            current_line = []
            for word in words:
                test_line = ' '.join(current_line + [word])
                w = fitz.get_text_length(test_line, fontname=font_name, fontsize=font_size)
                if w <= max_w:
                    current_line.append(word)
                else:
                    if current_line:
                        wrapped.append(' '.join(current_line))
                        current_line = [word]
                    else:
                        wrapped.append(word)
            if current_line:
                wrapped.append(' '.join(current_line))
            return wrapped

        for line in lines:
            if not line.strip():
                if current_line_on_page <= max_lines_per_page:
                    y += line_height
                    current_line_on_page += 1
                continue

            # Format prefixes & detect line type
            upper_line = line.strip()
            is_q = False
            is_a = False
            is_speaker = False

            if upper_line.upper().startswith('Q:') or upper_line.upper().startswith('Q.'):
                is_q = True
                line = 'Q.  ' + upper_line[2:].lstrip()
            elif upper_line.upper().startswith('A:') or upper_line.upper().startswith('A.'):
                is_a = True
                line = 'A.  ' + upper_line[2:].lstrip()
            elif ":" in upper_line and upper_line.split(":")[0].isupper() and len(upper_line.split(":")[0]) < 35:
                is_speaker = True

            # Identify headers and format them ALL CAPS
            is_header = any(line.startswith(prefix) for prefix in (
                'Task Name:', 'Document Title:', 'Version:', 'Status:', 'Generated At:', '--- '
            ))
            if is_header:
                line = line.upper()

            wrapped = get_wrapped_lines(line, max_width)
            for i, w_line in enumerate(wrapped):
                if current_line_on_page > max_lines_per_page:
                    page = start_new_page()
                    current_line_on_page = 1
                    y = top_margin
                
                # First line gets specific indent if Q/A/Speaker, wrapped continuation lines get text_start_x
                if i == 0:
                    if is_q:
                        indent_x = q_start_x
                    elif is_a:
                        indent_x = a_start_x
                    elif is_speaker:
                        indent_x = speaker_start_x
                    else:
                        indent_x = text_start_x
                else:
                    indent_x = text_start_x

                page.insert_text(fitz.Point(indent_x, y), w_line, fontsize=font_size, fontname=font_name)
                y += line_height
                current_line_on_page += 1

        pdf_bytes = doc.write()
        doc.close()
        return pdf_bytes

    async def get_document_word(self, task_id: UUID, org_id: UUID) -> bytes:
        """Get the generated AI document as a Word DOCX file."""
        from app.models.pipeline import PipelineRun

        task = await self.task_repo.get_task_in_org(task_id, org_id)
        if not task:
            raise NotFoundError("Task", str(task_id))

        doc = await self.pipeline_repo.get_document(task_id)
        if not doc:
            raise NotFoundError("AI Document")

        lines = []

        chunks = doc.corrected_chunks
        if not chunks:
            from app.models.transcript import Transcript
            tr_result = await self.db.execute(
                select(Transcript).where(Transcript.task_id == task_id)
            )
            transcript = tr_result.scalar_one_or_none()
            if transcript and transcript.matches:
                chunks = []
                for m in transcript.matches:
                    chunks.append({
                        "raw_chunk_id": m.get("raw_chunk_id"),
                        "original_raw_text": m.get("raw_chunk_text"),
                        "corrected_text": m.get("raw_chunk_text"),
                    })

        if chunks:
            for chunk in chunks:
                chunk_text = chunk.get("corrected_text") or chunk.get("original_raw_text") or ""
                lines.extend(self._split_qa_lines(chunk_text))
        else:
            for line in doc.content.split("\n"):
                lines.extend(self._split_qa_lines(line))

        # Check for cover page file
        import os
        from app.config import settings
        cover_path = None
        task_files = await self.task_repo.get_files(task_id)
        cover_files = [f for f in task_files if f.file_name.endswith("_cover.pdf")]
        if cover_files:
            cp = os.path.join(settings.STORAGE_PATH, cover_files[0].file_path)
            if os.path.exists(cp):
                cover_path = cp

        return self._generate_word_from_lines(doc.title, lines, cover_path)

    async def get_document_word_tracked(self, task_id: UUID, org_id: UUID) -> bytes:
        """
        Get the AI document as a Word DOCX with tracked changes.
        Deleted (original) text appears as red strikethrough; inserted (AI) text as blue underline.
        The reviewer can Accept / Reject each change individually in Word.
        """
        task = await self.task_repo.get_task_in_org(task_id, org_id)
        if not task:
            raise NotFoundError("Task", str(task_id))

        doc = await self.pipeline_repo.get_document(task_id)
        if not doc:
            raise NotFoundError("AI Document")

        # Build list of (original_text, corrected_text) pairs per Q/A block
        chunk_pairs: list[tuple[str, str]] = []

        chunks = doc.corrected_chunks or []
        if not chunks:
            from app.models.transcript import Transcript
            tr_result = await self.db.execute(
                select(Transcript).where(Transcript.task_id == task_id)
            )
            transcript = tr_result.scalar_one_or_none()
            if transcript and transcript.matches:
                for m in transcript.matches:
                    orig = m.get("raw_chunk_text") or ""
                    corr = m.get("raw_chunk_text") or ""
                    chunk_pairs.append((orig, corr))

        if not chunk_pairs:
            for chunk in chunks:
                orig = chunk.get("original_raw_text") or chunk.get("corrected_text") or ""
                corr = chunk.get("corrected_text") or chunk.get("original_raw_text") or ""
                chunk_pairs.append((orig, corr))

        # Check for cover page
        import os
        from app.config import settings
        cover_path = None
        task_files = await self.task_repo.get_files(task_id)
        cover_files = [f for f in task_files if f.file_name.endswith("_cover.pdf")]
        if cover_files:
            cp = os.path.join(settings.STORAGE_PATH, cover_files[0].file_path)
            if os.path.exists(cp):
                cover_path = cp

        return self._generate_word_tracked_changes(doc.title, chunk_pairs, cover_path)

    def _generate_word_tracked_changes(
        self,
        title: str,
        chunk_pairs: list[tuple[str, str]],
        cover_path: str = None
    ) -> bytes:
        """
        Generate a DOCX where every AI correction is shown as a tracked change:
          - Original text  → red strikethrough  (w:del / w:delText)
          - Corrected text → blue underline      (w:ins / w:t)

        The reviewer opens this in Word and uses Accept/Reject Changes to approve edits.
        Word-diff is performed at the word level for clean, readable diffs.
        """
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        import difflib, datetime, re, fitz

        document = docx.Document()

        # ── Page setup ──────────────────────────────────────────────────────
        section = document.sections[0]
        section.page_width  = int(8.5 * 914400)
        section.page_height = int(11  * 914400)
        section.left_margin   = int(1.25 * 914400)
        section.right_margin  = int(1.0  * 914400)
        section.top_margin    = int(1.0  * 914400)
        section.bottom_margin = int(1.0  * 914400)

        # ── Base paragraph style ─────────────────────────────────────────────
        style = document.styles["Normal"]
        font  = style.font
        font.name = "Courier New"
        font.size = Pt(12)

        # ── Author / date used in revision XML ───────────────────────────────
        AUTHOR = "Vertexlab AI"
        DATE   = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        _rev_id = [1]

        def _next_id():
            _rev_id[0] += 1
            return str(_rev_id[0])

        # ── XML helpers ──────────────────────────────────────────────────────
        def _make_rpr(bold=False, color: str = None, strike=False, underline=False):
            """Build a <w:rPr> element with font + optional colour/strikethrough/underline."""
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rfonts.set(qn(attr), "Courier New")
            rpr.append(rfonts)
            sz = OxmlElement("w:sz");  sz.set(qn("w:val"), "24"); rpr.append(sz)
            szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "24"); rpr.append(szCs)
            if bold:
                rpr.append(OxmlElement("w:b"))
            if color:
                col = OxmlElement("w:color"); col.set(qn("w:val"), color); rpr.append(col)
            if strike:
                rpr.append(OxmlElement("w:strike"))
            if underline:
                ul = OxmlElement("w:u"); ul.set(qn("w:val"), "single"); rpr.append(ul)
            return rpr

        def _ins_run(para_xml, text: str, bold=False):
            """Append a <w:ins><w:r><w:t> — blue underlined insertion."""
            ins = OxmlElement("w:ins")
            ins.set(qn("w:id"), _next_id())
            ins.set(qn("w:author"), AUTHOR)
            ins.set(qn("w:date"), DATE)
            r = OxmlElement("w:r")
            r.append(_make_rpr(bold=bold, color="1F4E79", underline=True))
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
            ins.append(r)
            para_xml.append(ins)

        def _del_run(para_xml, text: str, bold=False):
            """Append a <w:del><w:r><w:delText> — red strikethrough deletion."""
            d = OxmlElement("w:del")
            d.set(qn("w:id"), _next_id())
            d.set(qn("w:author"), AUTHOR)
            d.set(qn("w:date"), DATE)
            r = OxmlElement("w:r")
            r.append(_make_rpr(bold=bold, color="C00000", strike=True))
            dt = OxmlElement("w:delText")
            dt.set(qn("xml:space"), "preserve")
            dt.text = text
            r.append(dt)
            d.append(r)
            para_xml.append(d)

        def _plain_run(para_xml, text: str, bold=False):
            """Append a plain unchanged <w:r><w:t> run."""
            r = OxmlElement("w:r")
            r.append(_make_rpr(bold=bold))
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
            para_xml.append(r)

        def _new_para(doc_obj, ppr_xml=None):
            """Add a new paragraph and return its _p XML element."""
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = Pt(14)
            if ppr_xml is not None:
                p._p.insert(0, ppr_xml)
            return p._p

        def _ppr(left_indent_twips: int = 0):
            pPr = OxmlElement("w:pPr")
            if left_indent_twips:
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), str(left_indent_twips))
                pPr.append(ind)
            return pPr

        # ── Word-level diff between two strings ──────────────────────────────
        def _word_diff(orig: str, corr: str):
            """
            Tokenise both strings into words+spaces and return a list of
            ('equal'|'delete'|'insert'|'replace', orig_tokens, corr_tokens).
            """
            def tokenise(s):
                # Split keeping spaces so we can reconstruct readable text
                return re.findall(r'\S+|\s+', s)

            orig_toks = tokenise(orig)
            corr_toks = tokenise(corr)
            sm = difflib.SequenceMatcher(None, orig_toks, corr_toks, autojunk=False)
            ops = []
            for tag, i1, i2, j1, j2 in sm.get_opcodes():
                ops.append((tag, orig_toks[i1:i2], corr_toks[j1:j2]))
            return ops

        # ── Detect Q/A prefix ────────────────────────────────────────────────
        _qa_re = re.compile(r'^(Q|A)[\.:\s]\s*', re.IGNORECASE)

        def _split_prefix(text: str):
            """Return (prefix, body) e.g. ('Q. ', 'Good afternoon...')"""
            m = _qa_re.match(text.strip())
            if m:
                ltr = m.group(1).upper()
                body = text.strip()[m.end():]
                return f"{ltr}. ", body
            return None, text.strip()

        # ── Cover page ──────────────────────────────────────────────────────
        if cover_path:
            try:
                cover_doc = fitz.open(cover_path)
                for pg_idx in range(len(cover_doc)):
                    pg = cover_doc[pg_idx]
                    for block in pg.get_text("blocks"):
                        txt = block[4].strip()
                        if txt:
                            for ln in txt.split("\n"):
                                p = _new_para(document)
                                _plain_run(p, ln)
                cover_doc.close()
                document.add_page_break()
            except Exception:
                pass

        # ── Title ────────────────────────────────────────────────────────────
        tp = document.add_paragraph()
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(12)
        tr = tp.add_run(f"AI CORRECTIONS — {title or 'TRANSCRIPT'}")
        tr.bold = True
        tr.font.name = "Courier New"
        tr.font.size = Pt(12)

        # ── Legend ───────────────────────────────────────────────────────────
        leg_p = document.add_paragraph()
        leg_p.paragraph_format.space_after = Pt(10)
        leg_run = leg_p.add_run(
            "Legend:  "
        )
        leg_run.font.name = "Courier New"
        leg_run.font.size = Pt(10)

        del_run_leg = leg_p.add_run("Deleted text (original)  ")
        del_run_leg.font.name  = "Courier New"
        del_run_leg.font.size  = Pt(10)
        del_run_leg.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        del_run_leg.font.strike = True

        ins_run_leg = leg_p.add_run("Inserted text (AI correction)")
        ins_run_leg.font.name   = "Courier New"
        ins_run_leg.font.size   = Pt(10)
        ins_run_leg.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        ins_run_leg.font.underline = True

        document.add_paragraph()  # spacer

        # ── Body: one paragraph per chunk ────────────────────────────────────
        for orig_text, corr_text in chunk_pairs:
            orig_text = (orig_text or "").strip()
            corr_text = (corr_text or "").strip()
            if not orig_text and not corr_text:
                continue

            orig_prefix, orig_body = _split_prefix(orig_text)
            corr_prefix, corr_body = _split_prefix(corr_text)

            # Use corrected prefix (or orig if no change)
            prefix = corr_prefix or orig_prefix
            is_qa  = prefix is not None
            indent = 720 if is_qa else 0  # ~0.5 inch for Q/A body

            p = _new_para(document, _ppr(indent))

            # Prefix label (Q. / A.) — plain, bold
            if prefix:
                _plain_run(p, prefix, bold=True)

            # Word-level diff of the bodies
            ops = _word_diff(orig_body, corr_body)
            for tag, orig_toks, corr_toks in ops:
                orig_seg = "".join(orig_toks)
                corr_seg = "".join(corr_toks)
                if tag == "equal":
                    _plain_run(p, orig_seg)
                elif tag == "delete":
                    _del_run(p, orig_seg)
                elif tag == "insert":
                    _ins_run(p, corr_seg)
                elif tag == "replace":
                    _del_run(p, orig_seg)
                    _ins_run(p, corr_seg)

        # ── Serialize ────────────────────────────────────────────────────────
        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()



    def _generate_word_from_lines(self, title: str, lines: list[str], cover_path: str = None) -> bytes:
        """
        Generate a deposition-grade DOCX from transcript lines.

        Architecture:
        - Custom 'Transcript' style (never overwrites Normal).
        - Grouped paragraphs: each Q/A block is ONE <w:p> with <w:br/> between
          wrapped continuation lines. New <w:p> only on blank lines or Q/A transitions.
        - Full font specification (ascii/hAnsi/eastAsia/cs).
        - Raw XML injection for properties python-docx doesn't expose.
        - Word compatibility flags for consistent rendering.
        - Proper PAGE field code structure (begin → instrText → separate → fallback → end).
        """
        import docx
        from docx.shared import Pt, Inches, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT, WD_SECTION_START
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        import fitz
        import re

        document = docx.Document()

        # ── Helper: inject proper PAGE field code ──────────────────────────
        def add_page_number(run):
            """Insert PAGE field with begin → instrText → separate → fallback → end."""
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = " PAGE "

            fldChar_separate = OxmlElement('w:fldChar')
            fldChar_separate.set(qn('w:fldCharType'), 'separate')

            # Fallback display text (shown before field is updated)
            fallback_t = OxmlElement('w:t')
            fallback_t.text = "1"

            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar_begin)
            run._r.append(instrText)
            run._r.append(fldChar_separate)
            run._r.append(fallback_t)
            run._r.append(fldChar_end)

        # ── Helper: inject a boolean XML flag element ──────────────────────
        def _make_flag(tag: str, val: str = None) -> OxmlElement:
            """Create a simple OxmlElement, optionally with w:val attribute."""
            el = OxmlElement(tag)
            if val is not None:
                el.set(qn('w:val'), val)
            return el

        # ── Helper: classify a line as Q, A, or continuation ──────────────
        def _get_qa_prefix(line: str) -> str | None:
            """Return 'Q' or 'A' if the line starts with a Q/A indicator, else None."""
            stripped = line.strip().upper()
            if stripped.startswith('Q:') or stripped.startswith('Q.') or stripped.startswith('Q '):
                return 'Q'
            if stripped.startswith('A:') or stripped.startswith('A.') or stripped.startswith('A '):
                return 'A'
            return None

        # ── Helper: normalize Q/A prefix to standard form ─────────────────
        def _normalize_line(line: str) -> str:
            """Normalize Q:/A: and Q /A  prefixes to 'Q. '/'A. ' format."""
            stripped = line.strip()
            upper = stripped.upper()
            if upper.startswith('Q:') or upper.startswith('Q.') or upper.startswith('Q '):
                return 'Q. ' + stripped[2:].lstrip()
            if upper.startswith('A:') or upper.startswith('A.') or upper.startswith('A '):
                return 'A. ' + stripped[2:].lstrip()
            return line

        # ══════════════════════════════════════════════════════════════════
        # 1. COVER PAGE SECTION
        # ══════════════════════════════════════════════════════════════════
        cover_page_count = 0
        if cover_path:
            try:
                cover_doc = fitz.open(cover_path)
                cover_page_count = len(cover_doc)
                cover_section = document.sections[0]
                cover_section.page_width = Inches(8.5)
                cover_section.page_height = Inches(11.0)
                cover_section.left_margin = Inches(0)
                cover_section.right_margin = Inches(0)
                cover_section.top_margin = Inches(0)
                cover_section.bottom_margin = Inches(0)
                cover_section.header_distance = Inches(0)
                cover_section.footer_distance = Inches(0)
                cover_section.orientation = WD_ORIENT.PORTRAIT

                for i, page in enumerate(cover_doc):
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    p = document.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(BytesIO(img_data), width=Inches(8.5), height=Inches(11.0))
                    if i < len(cover_doc) - 1:
                        document.add_page_break()

                cover_doc.close()
                body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
            except Exception as e:
                logger.warning(f"Failed to merge cover PDF into Word document: {e}")
                cover_page_count = 0
                body_section = document.sections[0]
        else:
            body_section = document.sections[0]

        # ══════════════════════════════════════════════════════════════════
        # 2. BODY SECTION CONFIGURATION
        # ══════════════════════════════════════════════════════════════════
        body_section.page_width = Inches(8.5)
        body_section.page_height = Inches(11.0)
        body_section.left_margin = Inches(1.1)
        body_section.right_margin = Inches(1.1)
        body_section.top_margin = Inches(1.4)
        body_section.bottom_margin = Inches(1.0)
        body_section.orientation = WD_ORIENT.PORTRAIT
        body_section.header_distance = Inches(0.65)
        body_section.footer_distance = Inches(0.5)

        # Page numbering: continue from where the cover left off
        # e.g. 4-page cover → body starts at page 5
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), str(cover_page_count + 1))
        body_section._sectPr.append(pgNumType)

        # Line Numbers: restart each page, 0.25" from text
        sectPr = body_section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:start'), '1')
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:restart'), 'newPage')
        lnNumType.set(qn('w:distance'), '360')  # 0.25 inches = 360 DXA
        sectPr.append(lnNumType)

        # Document grid: linePitch for 25 lines per page
        # Available height = 11" - 1.4" - 1.0" = 8.6" = 619.2pt = 12384 twips
        # linePitch = 12384 / 25 ≈ 495 twips
        docGrid = OxmlElement('w:docGrid')
        docGrid.set(qn('w:type'), 'lines')
        docGrid.set(qn('w:linePitch'), '495')
        sectPr.append(docGrid)

        # Page Number in Header (Top-Right)
        header = body_section.header
        if cover_path:
            header.is_linked_to_previous = False
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run()
        # Set font on header run via XML for full coverage
        header_rPr = header_run._r.get_or_add_rPr()
        h_rFonts = OxmlElement('w:rFonts')
        h_rFonts.set(qn('w:ascii'), 'Courier New')
        h_rFonts.set(qn('w:hAnsi'), 'Courier New')
        h_rFonts.set(qn('w:eastAsia'), 'Courier New')
        h_rFonts.set(qn('w:cs'), 'Courier New')
        header_rPr.append(h_rFonts)
        h_sz = OxmlElement('w:sz')
        h_sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
        header_rPr.append(h_sz)
        h_szCs = OxmlElement('w:szCs')
        h_szCs.set(qn('w:val'), '20')
        header_rPr.append(h_szCs)
        add_page_number(header_run)

        # ══════════════════════════════════════════════════════════════════
        # 3. CREATE CUSTOM 'Transcript' STYLE (do NOT overwrite Normal)
        # ══════════════════════════════════════════════════════════════════
        transcript_style = document.styles.add_style('Transcript', WD_STYLE_TYPE.PARAGRAPH)
        transcript_style.base_style = document.styles['Normal']

        # ── 3a. Full font specification via raw XML ───────────────────────
        style_rPr = transcript_style.element.get_or_add_rPr()

        # Font family: all four slots
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:eastAsia'), 'Courier New')
        rFonts.set(qn('w:cs'), 'Courier New')
        style_rPr.append(rFonts)

        # Font size: 10pt (20 half-points) for both ascii and cs
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')
        style_rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '20')
        style_rPr.append(szCs)

        # Character scaling: 100%
        w_elem = OxmlElement('w:w')
        w_elem.set(qn('w:val'), '100')
        style_rPr.append(w_elem)

        # Character spacing: Normal (0)
        spacing_r = OxmlElement('w:spacing')
        spacing_r.set(qn('w:val'), '0')
        style_rPr.append(spacing_r)

        # Kerning: Off
        kern = OxmlElement('w:kern')
        kern.set(qn('w:val'), '0')
        style_rPr.append(kern)

        # Position: Normal baseline
        position = OxmlElement('w:position')
        position.set(qn('w:val'), '0')
        style_rPr.append(position)

        # Language
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'en-US')
        lang.set(qn('w:eastAsia'), 'en-US')
        lang.set(qn('w:bidi'), 'ar-SA')
        style_rPr.append(lang)

        # ── 3b. Paragraph format via python-docx API ──────────────────────
        pf = transcript_style.paragraph_format
        pf.line_spacing = 1.0           # Single
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent = Inches(0)
        pf.right_indent = Inches(0)
        pf.first_line_indent = Inches(0)
        pf.widow_control = False
        pf.keep_with_next = False
        pf.keep_together = False
        pf.page_break_before = False

        # ── 3c. Paragraph properties only accessible via raw XML ──────────
        style_pPr = transcript_style.element.get_or_add_pPr()

        # Don't auto-remove spacing between same-style paragraphs
        style_pPr.append(_make_flag('w:contextualSpacing', '0'))
        # Don't snap to document grid (critical for fixed-pitch transcripts)
        style_pPr.append(_make_flag('w:snapToGrid', '0'))
        # Don't auto-space between East Asian and Latin characters
        style_pPr.append(_make_flag('w:autoSpaceDE', '0'))
        # Don't auto-space between numbers and Latin characters
        style_pPr.append(_make_flag('w:autoSpaceDN', '0'))
        # Don't auto-adjust right indent
        style_pPr.append(_make_flag('w:adjustRightInd', '0'))
        # Never hyphenate transcript text
        style_pPr.append(_make_flag('w:suppressAutoHyphens'))
        # Baseline text alignment (standard for monospace)
        style_pPr.append(_make_flag('w:textAlignment', 'baseline'))
        # Show line numbers (do not suppress)
        style_pPr.append(_make_flag('w:suppressLineNumbers', '0'))

        # ══════════════════════════════════════════════════════════════════
        # 4. WORD COMPATIBILITY FLAGS
        # ══════════════════════════════════════════════════════════════════
        settings_element = document.settings.element
        compat = OxmlElement('w:compat')
        compat_flags = [
            'usePrinterMetrics',
            'doNotExpandShiftReturn',
            'ulTrailSpace',
            'doNotUseHTMLParagraphAutoSpacing',
            'layoutRawTableWidth',
            'doNotBreakWrappedTables',
            'doNotSnapToGridInCell',
            'dontWrapTextWithPunct',
            'dontUseEastAsianBreakRules',
        ]
        for flag in compat_flags:
            compat.append(OxmlElement(f'w:{flag}'))
        settings_element.append(compat)

        # ══════════════════════════════════════════════════════════════════
        # 5. BUILD TRANSCRIPT BODY — GROUPED PARAGRAPH MODEL
        # ══════════════════════════════════════════════════════════════════
        # Strategy:
        #   - Each Q block → 1 <w:p>, each A block → 1 <w:p>
        #   - Continuation lines within a block use <w:br/> (line break), not new <w:p>
        #   - Blank lines → flush current buffer as a paragraph, emit empty <w:p>
        #   - Q/A prefix transition → flush current buffer, start new buffer

        def _flush_buffer(buf: list[str], doc_obj):
            """Emit a single <w:p> from accumulated lines, joined by <w:br/>."""
            if not buf:
                return
            p = doc_obj.add_paragraph(style='Transcript')

            for idx, text in enumerate(buf):
                if idx > 0:
                    # Insert line break <w:br/> before continuation lines
                    br_run = p.add_run()
                    br_elem = OxmlElement('w:br')
                    br_run._r.append(br_elem)

                # Determine if this line's Q./A. prefix should be bold
                is_qa = bool(re.match(r'^[QA]\.\s', text))
                if is_qa:
                    # Bold the "Q. " or "A. " prefix, normal weight for the rest
                    prefix = text[:3]  # "Q. " or "A. "
                    rest = text[3:]
                    prefix_run = p.add_run(prefix)
                    prefix_run.bold = True
                    if rest:
                        p.add_run(rest)
                else:
                    p.add_run(text)

        # Pre-process: normalize all Q/A prefixes
        normalized_lines = [_normalize_line(l) for l in lines]

        # Start the document body with "EXAMINATION" heading
        exam_p = document.add_paragraph(style='Transcript')
        exam_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        exam_run = exam_p.add_run('EXAMINATION')
        exam_run.bold = True

        buffer: list[str] = []
        current_prefix: str | None = None  # Track current Q/A block

        for line in normalized_lines:
            # Blank line → flush + emit empty paragraph
            if not line.strip():
                _flush_buffer(buffer, document)
                buffer = []
                current_prefix = None
                # Emit blank paragraph (spacer)
                document.add_paragraph(style='Transcript')
                continue

            # Skip metadata headers (they should have been removed upstream,
            # but guard here as well)
            is_header = any(line.startswith(prefix) for prefix in (
                'Task Name:', 'Document Title:', 'Version:',
                'Status:', 'Generated At:', '--- '
            ))
            if is_header:
                continue

            # Detect Q/A prefix
            line_prefix = _get_qa_prefix(line)

            if line_prefix is not None:
                # This line starts a new Q or A block
                if line_prefix != current_prefix:
                    # Prefix changed (Q→A, A→Q, or None→Q/A) → flush old buffer
                    _flush_buffer(buffer, document)
                    buffer = []
                    current_prefix = line_prefix
                # Add to current block
                buffer.append(line)
            else:
                # Continuation line (no Q/A prefix) — append to current buffer
                buffer.append(line)

        # Flush any remaining lines
        _flush_buffer(buffer, document)

        # ══════════════════════════════════════════════════════════════════
        # 6. SAVE & RETURN
        # ══════════════════════════════════════════════════════════════════
        file_stream = BytesIO()
        document.save(file_stream)
        return file_stream.getvalue()


def _format_time(seconds: float) -> str:
    """Format seconds into HH:MM:SS format."""
    if seconds is None:
        return "00:00:00"
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"

